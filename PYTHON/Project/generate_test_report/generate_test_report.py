from docx import Document
from datetime import datetime
from tkinter import Tk, filedialog

def update_report_template(template_path):
    # 打开现有的Word文档
    doc = Document(template_path)

     # 遍历文档中的段落
    for paragraph in doc.paragraphs:
        if '測試人(Tester)' in paragraph.text:
            paragraph.text=paragraph.text.replace('中文名', '周志鵬')
        if '製表日期(Date)' in paragraph.text:
            # 在"製表日期(Date):"后面添加当前日期
            date_str = datetime.now().strftime('%Y/%m/%d')
            paragraph.text=paragraph.text.replace('yyyy/mm/dd', date_str)

    table=doc.tables[0]
    table.cell(0, 1).text = 'PRDG010'
    table.cell(0, 3).text = 'PRDG010_DESC'
    table.cell(1, 3).text = 'MOD_DESC'
    table.cell(2, 1).text = 'I-1234567'
    # 保存修改后的文档
    doc.save('modified_report_with_table.docx')
if __name__ == "__main__":
    root = Tk()
    root.withdraw()  # 隐藏主窗口

    # 打开文件对话框，选择Word文档
    file_path = filedialog.askopenfilename(title="选择Word文档", filetypes=[("Word文档", "*.docx")])

    if file_path:
        # 替换为你的测试人名字
        tester_name = 'XXX'
        update_report_template(file_path)
        print("处理完成！")
